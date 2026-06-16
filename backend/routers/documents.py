from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from fastapi.responses import FileResponse, StreamingResponse
from typing import List, Optional
import io
import json
import urllib.parse
from docx import Document as DocxDocument

from database import get_db
import models, schemas
from auth_utils import get_current_user

router = APIRouter(prefix="/api/documents", tags=["Документы"])

DOC_PREFIXES = {
    "vacation":   "ОТП",
    "travel":     "КМД",
    "waybill":    "ТМ",
    "employment": "ТД",
    "transport_waybill": "ТРН",
}

STATUS_LABELS = {
    "draft":    "Черновик",
    "pending":  "На рассмотрении",
    "approved": "Одобрено",
    "rejected": "Отклонено",
}

TYPE_LABELS = {
    "vacation":   "Заявка на отпуск",
    "travel":     "Командировка",
    "waybill":    "Путевой лист",
    "employment": "Трудовой договор",
    "transport_waybill": "Транспортная накладная",
}


def _enrich(doc: models.Document) -> schemas.DocumentResponse:
    r = schemas.DocumentResponse.model_validate(doc)
    
    # Check if doc has a related created_by_user and what name to use
    if doc.created_by_user:
        # User model may have `full_name` or fall back to `name` if person model is used directly
        r.created_by_name = getattr(doc.created_by_user, 'full_name', getattr(doc.created_by_user, 'name', '—'))
    else:
        r.created_by_name = "—"
        
    return r


@router.get("/", response_model=List[schemas.DocumentResponse])
def list_documents(
    doc_type: Optional[str] = None,
    db: Session = Depends(get_db),
    _: models.User = Depends(get_current_user)
):
    q = db.query(models.Document)
    if doc_type:
        q = q.filter(models.Document.document_type == doc_type)
    return [_enrich(d) for d in q.order_by(models.Document.created_at.desc()).all()]

@router.get("/export_xml_employment")
def export_employment_xml(db: Session = Depends(get_db)):
    from fastapi import Response
    import json, urllib.parse
    docs = db.query(models.Document).filter(models.Document.document_type == 'employment').all()
    
    lines = ['<?xml version="1.0" encoding="UTF-8"?>', '<ТрудовыеДоговоры>']
    
    for doc in docs:
        extra = {}
        if doc.extra_data:
            try:
                extra = json.loads(doc.extra_data)
            except:
                pass
        
        def v(key, default=''):
            val = extra.get(key, default)
            return str(val) if val is not None else ''
        
        # Булевы значения
        def b(key):
            val = extra.get(key)
            return 'true' if val in (True, 'true', '1', 1) else 'false'
        
        lines.append('\t<ТрудовойДоговор>')
        # --- Главное ---
        lines.append(f'\t\t<Организация>{v("organization") or "ООО «НОРД-ИСТ ГРУПП»"}</Организация>')
        lines.append(f'\t\t<Сотрудник>{doc.employee_name}</Сотрудник>')
        lines.append(f'\t\t<Дата>{doc.created_at.strftime("%Y-%m-%d")}</Дата>')
        lines.append(f'\t\t<Номер>{v("contract_number") or doc.number}</Номер>')
        lines.append(f'\t\t<ВидДоговора>{v("contract_type")}</ВидДоговора>')
        lines.append(f'\t\t<ДатаПриема>{v("hire_date")}</ДатаПриема>')
        lines.append(f'\t\t<Подразделение>{v("department")}</Подразделение>')
        lines.append(f'\t\t<Территория>{v("territory")}</Территория>')
        lines.append(f'\t\t<ДолжностьПоШтату>{v("position_staff")}</ДолжностьПоШтату>')
        lines.append(f'\t\t<Должность>{v("job_title")}</Должность>')
        lines.append(f'\t\t<ГрафикРаботы>{v("work_schedule")}</ГрафикРаботы>')
        lines.append(f'\t\t<ВидЗанятости>{v("employment_type")}</ВидЗанятости>')
        # --- Второстепенное ---
        lines.append(f'\t\t<ОтразитьВТрудовойКнижке>{b("reflect_in_workbook")}</ОтразитьВТрудовойКнижке>')
        lines.append(f'\t\t<ТрудоваяФиксация>{v("work_fixation")}</ТрудоваяФиксация>')
        lines.append(f'\t\t<НаименованиеДокумента>{v("doc_name")}</НаименованиеДокумента>')
        lines.append(f'\t\t<НачалоТрудовойДеятельности>{b("start_of_work")}</НачалоТрудовойДеятельности>')
        lines.append(f'\t\t<СпособВедения>{v("management_method")}</СпособВедения>')
        lines.append(f'\t\t<ДатаЗаявленияОВыбореСпоспоба>{v("method_choice_date")}</ДатаЗаявленияОВыбореСпоспоба>')
        lines.append(f'\t\t<НаименованиеВторДок>{v("second_doc_name")}</НаименованиеВторДок>')
        lines.append(f'\t\t<ДатаВторогоДок>{v("second_doc_date")}</ДатаВторогоДок>')
        lines.append(f'\t\t<СерияВторогоДок>{v("second_doc_series")}</СерияВторогоДок>')
        lines.append(f'\t\t<НомерВторогоДок>{v("second_doc_number")}</НомерВторогоДок>')
        lines.append(f'\t\t<ПКУ>{v("pku")}</ПКУ>')
        lines.append(f'\t\t<Разряд>{v("grade")}</Разряд>')
        lines.append(f'\t\t<ФОТ>{v("fot")}</ФОТ>')
        lines.append(f'\t\t<Ответственный>{v("responsible")}</Ответственный>')
        lines.append('\t</ТрудовойДоговор>')
    
    lines.append('</ТрудовыеДоговоры>')
    xml_str = '\n'.join(lines)
    
    # Сохраняем локально на Google Диск (резервная копия, если папка смонтирована)
    try:
        import os
        export_dir = r"G:\Мой диск\XML files"
        os.makedirs(export_dir, exist_ok=True)
        with open(os.path.join(export_dir, 'trudovye_dogovory.xml'), 'w', encoding='utf-8') as f:
            f.write(xml_str)
    except Exception as e:
        print(f"Предупреждение: Не удалось сохранить XML на локальный Google Диск: {str(e)}")

    # Отправляем напрямую в облако Google Drive через Google Apps Script Web App
    try:
        import httpx
        script_url = os.getenv("GOOGLE_SCRIPT_URL", "https://script.google.com/macros/s/AKfycbxQJK61xo9pIqb-wM1yIPxKTpMRokhu3_2SVGEKTP046GTXqZsV0B9WBxjfrs-9cBE/exec")
        response = httpx.post(
            script_url, 
            json={"filename": "trudovye_dogovory.xml", "content": xml_str}, 
            follow_redirects=True, 
            timeout=10.0
        )
        print(f"Отправка trudovye_dogovory.xml в Google Drive: {response.status_code}, {response.text}")
    except Exception as e:
        print(f"Предупреждение: Не удалось отправить XML в облако Google Drive: {str(e)}")
    
    filename = 'trudovye_dogovory.xml'
    filename_encoded = urllib.parse.quote(filename)
    headers = {
        "Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}",
        "Cache-Control": "no-cache, no-store, must-revalidate",
    }
    return Response(content=xml_str, media_type="application/xml; charset=utf-8", headers=headers)


@router.get("/export_xml")
def export_docs_xml(db: Session = Depends(get_db)):
    from fastapi import Response
    import json
    docs = db.query(models.Document).filter(models.Document.document_type == 'waybill').all()
    xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n<Документы>\n'
    
    for doc in docs:
        extra = {}
        if doc.extra_data:
            try:
                extra = json.loads(doc.extra_data)
            except:
                pass
        
        fuel_mark = extra.get('fuel_mark', '')
        if not fuel_mark:
            continue
            
        fuel_qty = float(extra.get('fuel_issued') or 0)
        fuel_cost = float(extra.get('fuel_cost') or 0)
        price = round(fuel_cost / fuel_qty, 2) if fuel_qty > 0 else fuel_cost

        xml_str += f'\t<ПутевойЛист>\n'
        xml_str += f'\t\t<ДатаСоздания>{doc.created_at.strftime("%Y-%m-%d")}</ДатаСоздания>\n'
        xml_str += f'\t\t<Номенклатура>{fuel_mark}</Номенклатура>\n'
        xml_str += f'\t\t<Количество>{fuel_qty}</Количество>\n'
        xml_str += f'\t\t<Цена>{price}</Цена>\n'
        xml_str += f'\t</ПутевойЛист>\n'
    xml_str += '</Документы>'
    
    # Сохраняем локально на Google Диск (резервная копия, если папка смонтирована)
    try:
        import os
        export_dir = r"G:\Мой диск\XML files"
        os.makedirs(export_dir, exist_ok=True)
        with open(os.path.join(export_dir, 'documents_waybills.xml'), 'w', encoding='utf-8') as f:
            f.write(xml_str)
    except Exception as e:
        print(f"Предупреждение: Не удалось сохранить XML на локальный Google Диск: {str(e)}")

    # Отправляем напрямую в облако Google Drive через Google Apps Script Web App
    try:
        import httpx
        script_url = os.getenv("GOOGLE_SCRIPT_URL", "https://script.google.com/macros/s/AKfycbxQJK61xo9pIqb-wM1yIPxKTpMRokhu3_2SVGEKTP046GTXqZsV0B9WBxjfrs-9cBE/exec")
        response = httpx.post(
            script_url, 
            json={"filename": "documents_waybills.xml", "content": xml_str}, 
            follow_redirects=True, 
            timeout=10.0
        )
        print(f"Отправка documents_waybills.xml в Google Drive: {response.status_code}, {response.text}")
    except Exception as e:
        print(f"Предупреждение: Не удалось отправить XML в облако Google Drive: {str(e)}")
        
    headers = {
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0"
    }
    return Response(content=xml_str, media_type="application/xml", headers=headers)

@router.post("/", response_model=schemas.DocumentResponse)
def create_document(
    body: schemas.DocumentCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    prefix = DOC_PREFIXES.get(body.document_type, "ДОК")
    count = db.query(models.Document).filter(
        models.Document.document_type == body.document_type
    ).count() + 1
    number = f"{prefix}-2026-{str(count).zfill(3)}"

    doc = models.Document(
        number=number,
        document_type=body.document_type,
        employee_name=body.employee_name,
        status="draft",
        created_by_id=current_user.id,
        extra_data=body.extra_data
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _enrich(doc)


@router.patch("/{doc_id}", response_model=schemas.DocumentResponse)
def update_status(
    doc_id: int,
    body: schemas.DocumentStatusUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")

    # Проверяем права в зависимости от нового статуса
    if body.status in ("approved", "rejected"):
        user_roles = [r.role_type for r in current_user.roles]
        if not any(role in user_roles for role in ("director", "manager", "admin")):
            raise HTTPException(status_code=403, detail="Только директор или менеджер могут согласовывать документы")
    elif body.status == "pending":
        if doc.status != "draft":
            raise HTTPException(status_code=400, detail="Только черновики можно отправлять на рассмотрение")
        # Любой авторизованный пользователь может перевести свой (или любой) черновик на рассмотрение

    # Сохраняем причину отказа если есть
    if body.reason:
        extra = {}
        if doc.extra_data:
            try:
                import json
                extra = json.loads(doc.extra_data)
            except:
                pass
        extra['rejection_reason'] = body.reason
        import json
        doc.extra_data = json.dumps(extra, ensure_ascii=False)

    doc.status = body.status
    db.commit()

    # Генерируем уведомление
    if body.status in ("approved", "rejected") and doc.created_by_id:
        msg = f"Документ {doc.number} был {'одобрен' if body.status == 'approved' else 'отклонен'}."
        if body.status == "rejected" and body.reason:
            msg += f" Причина: {body.reason}"
            
        notif = models.Notification(
            user_id=doc.created_by_id,
            title="Документ " + ("одобрен" if body.status == "approved" else "отклонен"),
            message=msg,
            icon="fas fa-check-circle" if body.status == "approved" else "fas fa-times-circle",
            tc="tg" if body.status == "approved" else "tr"
        )
        db.add(notif)
        db.commit()

    db.refresh(doc)
    return _enrich(doc)


@router.delete("/{doc_id}", status_code=204)
def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    user_roles = [r.role_type for r in current_user.roles]
    if not any(role in user_roles for role in ("director", "manager", "admin")):
        raise HTTPException(status_code=403, detail="Недостаточно прав")
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")
    db.delete(doc)
    db.commit()

@router.get("/{doc_id}/download")
def download_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")

    # Создаем Word-документ
    document = DocxDocument()
    
    extra = {}
    if doc.extra_data:
        try:
            extra = json.loads(doc.extra_data)
        except:
            pass

    if doc.document_type == 'waybill':
        document.add_heading(f"ПУТЕВОЙ ЛИСТ № {extra.get('series_number') or doc.number}", 0)
        document.add_paragraph("Организация: ООО «НОРД-ИСТ ГРУПП»")
        document.add_paragraph(f"Дата создания: {doc.created_at.strftime('%d.%m.%Y')}")
        document.add_paragraph(f"Ответственный (Сотрудник): {doc.employee_name}")
        document.add_paragraph(f"Транспортное средство: {extra.get('vehicle', '—')}")
        document.add_paragraph(f"Водитель: {extra.get('driver', '—')}")
        document.add_paragraph(f"Период действия: {extra.get('period_info', '—')}")
        
        document.add_heading("Задание водителю", level=1)
        table1 = document.add_table(rows=2, cols=2)
        table1.style = 'Table Grid'
        table1.cell(0, 0).text = "Откуда"
        table1.cell(0, 1).text = "Куда"
        table1.cell(1, 0).text = str(extra.get('route_from') or '—')
        table1.cell(1, 1).text = str(extra.get('route_to') or '—')
        
        document.add_paragraph()
        
        table2 = document.add_table(rows=2, cols=3)
        table2.style = 'Table Grid'
        table2.cell(0, 0).text = "Время выезда"
        table2.cell(0, 1).text = "Время возвращения"
        table2.cell(0, 2).text = "Дистанция (км)"
        dep_time = str(extra.get('departure_time') or '').replace('T', ' ')
        arr_time = str(extra.get('arrival_time') or '').replace('T', ' ')
        table2.cell(1, 0).text = dep_time if dep_time else '—'
        table2.cell(1, 1).text = arr_time if arr_time else '—'
        table2.cell(1, 2).text = str(extra.get('distance_km') or '—')

        document.add_heading("Движение горючего", level=1)
        table3 = document.add_table(rows=2, cols=3)
        table3.style = 'Table Grid'
        table3.cell(0, 0).text = "Марка топлива"
        table3.cell(0, 1).text = "Выдано"
        table3.cell(0, 2).text = "Остаток / Сдано"
        table3.cell(1, 0).text = str(extra.get('fuel_mark') or '—')
        table3.cell(1, 1).text = str(extra.get('fuel_issued') or '0')
        table3.cell(1, 2).text = str(extra.get('fuel_handed_over') or '0')
    else:
        document.add_heading(f"{TYPE_LABELS.get(doc.document_type, 'Документ')} № {doc.number}", 0)

        document.add_paragraph(f"Дата создания: {doc.created_at.strftime('%d.%m.%Y')}")
        document.add_paragraph(f"Сотрудник: {doc.employee_name}")
        creator_name = getattr(doc.created_by_user, 'full_name', getattr(doc.created_by_user, 'name', '—')) if doc.created_by_user else '—'
        document.add_paragraph(f"Составил: {creator_name}")
        document.add_paragraph(f"Статус: {STATUS_LABELS.get(doc.status, doc.status)}")
        
        # Добавляем дополнительные данные
        if doc.extra_data:
            document.add_heading("Детали", level=1)
            if extra:
                for k, v in extra.items():
                    if v:
                        document.add_paragraph(f"{str(k).upper()}: {v}")
            else:
                document.add_paragraph(doc.extra_data)
                
    document.add_paragraph()
    document.add_paragraph("______________________________")
    document.add_paragraph("Документ сформирован автоматически системой NortGru.")
    
    # Сохраняем в In-Memory буфер
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    # Форматируем имя файла, чтобы избежать проблем с русскими буквами
    filename = f"{doc.number}.docx"
    filename_encoded = urllib.parse.quote(filename)
    
    headers = {
        "Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"
    }

    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
